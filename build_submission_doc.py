from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from runtime_paths import EXPORT_ROOT, ensure_runtime_dirs


TEMPLATE = Path(r"D:\计算机思维\2026春人工智能通识与实践期末作品要求\作品说明文档模板.docx")
OUTPUT = EXPORT_ROOT / "古诗文全文查询_作品说明文档.docx"


def set_run_font(run, name: str, size: int, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def replace_paragraph(paragraph, text: str, *, align=None, size=12, bold=False):
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        run = paragraph.add_run(text)
    else:
        run = paragraph.runs[0]
        run.text = text
    set_run_font(run, "宋体", size, bold=bold)
    if align is not None:
        paragraph.alignment = align


def write_cell(cell, text: str, *, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    for idx, block in enumerate(text.split("\n")):
        if idx == 0:
            run = paragraph.add_run(block)
        else:
            run = paragraph.add_run()
            run.add_break()
            run = paragraph.add_run(block)
        set_run_font(run, "宋体", size, bold=bold)


def main():
    ensure_runtime_dirs()
    doc = Document(TEMPLATE)

    replace_paragraph(
        doc.paragraphs[0],
        "古诗文全文查询与拼音排版系统",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=18,
        bold=True,
    )

    summary_table = doc.tables[0]
    write_cell(
        summary_table.cell(0, 1),
        "本作品面向古诗文学习与课堂展示场景，解决传统检索方式中“查全文慢、朗读标注弱、展示排版杂”的问题。"
        "用户输入诗文题目后，系统可以自动检索全文，并以统一的古籍风格版式展示内容，降低教师备课和学生查阅成本。",
    )
    write_cell(
        summary_table.cell(1, 1),
        "作品基于本地网页应用实现，核心功能包括：\n"
        "1. 输入题目后自动查询古诗或古文全文；\n"
        "2. 按“拼音在上、汉字在下、逐字方格”的方式排版展示；\n"
        "3. 支持不同作品统一套用同一展示界面，而非单篇静态页面；\n"
        "4. 提供 PDF 效果预览与导出，便于课堂展示、打印和提交存档；\n"
        "5. 优化本地启动流程，保证查询页面可稳定打开。",
    )

    members_table = doc.tables[1]
    group_name = "古诗文智能检索与排版小组"
    for col in range(4):
        write_cell(members_table.cell(0, col), group_name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(members_table.cell(1, col), "小组成员", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        write_cell(members_table.cell(2, col), ["学号", "姓名", "班级", "完成内容"][col], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    rows = [
        ("202308083052", "彭天泰", "待补充", "需求梳理、界面参考整理、作品提交材料统筹。"),
        ("202407013065", "韩子轩", "24播音班", "查询页面搭建、本地启动脚本编写、页面交互联调。"),
        ("202407013066", "王裴", "24播音班", "古诗文查询逻辑接入、拼音排版实现、PDF 预览导出。"),
    ]
    for row_idx, row_data in enumerate(rows, start=3):
        for col_idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx < 3 else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(members_table.cell(row_idx, col_idx), value, align=align)

    log_table = doc.tables[2]
    logs = [
        ("第 1 轮", "Codex", "制作一个古诗文查询窗口，输入题目后能调出全文。", "初版依赖的 Python 启动路径错误，导致窗口无法正常打开。", "改为使用稳定运行时路径，并检查本地启动流程。"),
        ("第 2 轮", "Codex", "给查询结果增加拼音，并要求拼音在汉字上一行逐字对齐。", "最初只能输出普通文本，无法做到逐字格与拼音同步。", "重构前端渲染逻辑，改成统一的逐字排版组件。"),
        ("第 3 轮", "Codex", "复刻参考界面，而且不只是《岳阳楼记》，任意作品都要按同样方式呈现。", "静态界面只能展示单篇内容，短诗与古文在布局上也有差异。", "接入通用查询接口，统一标题、作者、正文和预览区的动态渲染。"),
        ("第 4 轮", "Codex", "完善 PDF 预览和导出，并修复本地页面偶发打不开的问题。", "出现过端口监听异常和页面报 ERR_CONNECTION_REFUSED 的情况。", "增加健康检查与独立后台启动，确保提交演示时可稳定访问。"),
    ]
    for row_idx, log in enumerate(logs, start=1):
        for col_idx, value in enumerate(log):
            align = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(log_table.cell(row_idx, col_idx), value, align=align, size=10)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
